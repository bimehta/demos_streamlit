import streamlit as st
import time
from uuid import uuid4
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pinecone import Pinecone, ServerlessSpec
from PyPDF2 import PdfReader
import docx
from langchain_openai import ChatOpenAI

# Pinecone and OpenAI API keys
pc = Pinecone(api_key="pcsk_2gtr65_Dm42ejX9QrrmBupkbKMP8PkbLm2fXCqiKKf8XPzW3gvZk6m86KEKPSGy4Lvzs9b")
st.title("Document Search with OpenAI API Key Input")

# Input OpenAI API Key
openai_api_key = st.text_input(
    "Enter your OpenAI API Key:",
    type="password",
    help="Your OpenAI API key is required to generate summaries and perform searches."
)

if not openai_api_key:
    st.warning("Please enter your OpenAI API key to proceed.")
    st.stop()

# index_name = "sooch-demo-index-deployed"

# # Initialize Pinecone index
# existing_indexes = [index_info["name"] for index_info in pc.list_indexes()]
# if index_name not in existing_indexes:
#     pc.create_index(
#         name=index_name,
#         dimension=3072,
#         metric="cosine",
#         spec=ServerlessSpec(cloud="aws", region="us-east-1"),
#     )
#     while not pc.describe_index(index_name).status["ready"]:
#         time.sleep(1)

# index = pc.Index(index_name)
# embeddings = OpenAIEmbeddings(model="text-embedding-3-large", openai_api_key=openai_api_key)
# vector_store = PineconeVectorStore(index=index, embedding=embeddings)

# Streamlit app
st.set_page_config(page_title="Sooch RAG Agent", page_icon="🤖")
st.image('SoochLogo.png', use_container_width=False) 
st.title("Sooch Document AI Q&A")


# Function to extract text from different file types
def extract_text(file):
    file_type = file.name.split('.')[-1].lower()
    if file_type == "txt":
        return file.read().decode("utf-8")
    elif file_type == "pdf":
        pdf_reader = PdfReader(file)
        extracted_text = []
        for page_num, page in enumerate(pdf_reader.pages, start=1):
            extracted_text.append({"text": page.extract_text(), "page": page_num})
        return extracted_text

    elif file_type == "docx":
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    else:
        st.error(f"Unsupported file type: {file_type}")
        return None
    
def chunk_text_with_metadata(text_with_meta, chunk_size=1000, overlap=200):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        length_function=len
    )
    chunks = []
    for entry in text_with_meta:
        text_chunks = text_splitter.split_text(entry["text"])
        for chunk in text_chunks:
            chunks.append({"chunk": chunk, "page": entry["page"]})
    return chunks

# Function to retrieve and summarize answers
def retrieve_and_summarize(query, k=3):
    # results = vector_store.similarity_search(query, k=k)
    results = []
    all_indexes = [index["name"] for index in pc.list_indexes()]  # Extract index names as strings
    for index_name in all_indexes:
        index = pc.Index(index_name)
        vector_store = PineconeVectorStore(index=index, embedding=OpenAIEmbeddings(model="text-embedding-3-large", openai_api_key=openai_api_key))
        try:
            # Perform similarity search
            index_results = vector_store.similarity_search(query, k=k)
            results.extend(index_results)
        except Exception as e:
            st.error(f"Error searching index {index_name}: {e}")
    if not results:
        return "No relevant documents found.", []

    # Combine all retrieved content for summarization
    combined_content = "\n".join([res.page_content for res in results])

    # Deduplicate sources by grouping chunks from the same document
    grouped_sources = {}
    for res in results:
        source = res.metadata.get("source", "Unknown")
        page = res.metadata.get("page", "N/A")
        link = res.metadata.get("link", "#")
        
        if source not in grouped_sources:
            grouped_sources[source] = {
                "link": link,
                "pages": set(),
                "snippets": []
            }
        grouped_sources[source]["pages"].add(page)
        grouped_sources[source]["snippets"].append(res.page_content[:200])  # Snippet from the chunk

    # Format sources with combined pages and snippets
    formatted_sources = []
    for source, data in grouped_sources.items():
        formatted_sources.append({
            "source": source,
            "link": data["link"],
            "pages": sorted(data["pages"]),
            "snippets": data["snippets"]
        })


    # Generate summarized answer
    # Using ChatOpenAI (gpt-3.5 or gpt-4) or plain OpenAI
    llm = ChatOpenAI(
        temperature=0,
        openai_api_key=openai_api_key,
        model_name="gpt-4o"  # or "gpt-4" if you have access
    )

    summary_prompt = (
        f"Summarize the following text to answer the query: '{query}'. Provide citations for each source:\n\n"
        f"{combined_content}\n\nAnswer:"
    )
    response = llm(summary_prompt)
    summary = response.content.strip()

    return summary, formatted_sources


# Upload documents
st.header("Upload your documents (PDF, DOCX, or TXT). You can upload multiple files.")
uploaded_files = st.file_uploader(
    "Upload your documents (PDF, DOCX, TXT)", accept_multiple_files=True, type=["txt", "pdf", "docx"]
)

# Check if index is empty
# index_stats = index.describe_index_stats()
# if index_stats["total_vector_count"] == 0:
#     st.warning("The vector storeis empty. Please upload documents to build the vector.")

# if uploaded_files:
#     if index_stats["total_vector_count"] == 0:
#         with st.spinner("Processing documents and updating vector store..."):
#             documents = []
#             for uploaded_file in uploaded_files:
#                 text_with_pages = extract_text(uploaded_file)
#                 if text_with_pages:
#                     chunks_with_metadata = chunk_text_with_metadata(text_with_pages)
#                     for i, chunk_meta in enumerate(chunks_with_metadata):
#                         doc = Document(
#                             page_content=chunk_meta["chunk"],
#                             metadata={
#                                 "source": f"{uploaded_file.name} - Chunk {i+1}",
#                                 "page": chunk_meta["page"],
#                                 "link": f"/path/to/{uploaded_file.name}#chunk-{i+1}"  # Placeholder for chunk link
#                             }
#                         )
#                         documents.append(doc)

#     # Generate unique IDs for documents and add them to the vector store
#         if documents:
#             uuids = [str(uuid4()) for _ in range(len(documents))]
#             vector_store.add_documents(documents=documents, ids=uuids)
#             st.success(f"Document Processed Succesfully")
#     else:
#         st.info("Documents already present in Vector Store. Skipping document processing.")

if uploaded_files:
    for uploaded_file in uploaded_files:
        # Use file name (without extension) as index name
        index_name = uploaded_file.name.split('.')[0].lower().replace(" ", "_")

        # Check if the index exists
        existing_indexes = [index_info["name"] for index_info in pc.list_indexes()]
        if index_name not in existing_indexes:
            pc.create_index(
                name=index_name,
                dimension=3072,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
            while not pc.describe_index(index_name).status["ready"]:
                time.sleep(1)

        index = pc.Index(index_name)
        vector_store = PineconeVectorStore(index=index, embedding=OpenAIEmbeddings(model="text-embedding-3-large", openai_api_key=openai_api_key))

        # Check if the index is empty
        index_stats = index.describe_index_stats()
        if index_stats["total_vector_count"] == 0:
            with st.spinner(f"Processing {uploaded_file.name} into vector store..."):
                text_with_pages = extract_text(uploaded_file)
                if text_with_pages:
                    chunks_with_metadata = chunk_text_with_metadata(text_with_pages)
                    documents = [
                        Document(
                            page_content=chunk_meta["chunk"],
                            metadata={
                                "source": f"{uploaded_file.name}",
                                "page": chunk_meta["page"],
                                "link": f"/path/to/{uploaded_file.name}#chunk-{i+1}"  # Placeholder for chunk link
                            }
                        )
                        for i, chunk_meta in enumerate(chunks_with_metadata)
                    ]

                    # Add documents to the vector store
                    uuids = [str(uuid4()) for _ in range(len(documents))]
                    vector_store.add_documents(documents=documents, ids=uuids)
                st.success(f"Processed {uploaded_file.name} into the index '{index_name}'.")
        else:
            st.info(f"The vector for document '{index_name}' already exists. Skipping processing for {uploaded_file.name}.")

# Question input
st.header("Ask the Sooch Agent a question about your uploaded documents")
question = st.text_input("Enter your question:")

if question:
    st.write("Searching and summarizing relevant information...")
    # Retrieve and summarize answer
    answer, sources = retrieve_and_summarize(question)

    # Display the results
    st.header("Sooch Agent Answer")
    st.write(answer)

    st.header("Source(s) of Information:")
    for source in sources:
        page_info = f"Pages: {', '.join(map(str, source['pages']))}" if source["pages"] else "Page: N/A"
        st.markdown(f"- [{source['source']}]({source['link']}) ({page_info})")
        for snippet in source["snippets"]:
            st.write(f"Snippet: {snippet}...")
